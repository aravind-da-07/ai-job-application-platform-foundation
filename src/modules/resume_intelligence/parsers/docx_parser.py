"""
DOCX Resume Parser.

Parses DOCX resumes using python-docx and returns a ResumeParseResult.
"""

from __future__ import annotations

import logging

from docx import Document

from src.modules.resume_intelligence.domain.exceptions.parsing import (
    InvalidResumeDocumentError,
)
from src.modules.resume_intelligence.parsers.base_parser import ResumeParser
from src.modules.resume_intelligence.schemas.resume_document import (
    ResumeDocument,
)
from src.modules.resume_intelligence.schemas.resume_parse_result import (
    ResumeParseResult,
)
from src.modules.resume_intelligence.utils.text_normalizer import (
    TextNormalizer,
)

logger = logging.getLogger(__name__)


class DOCXParser(ResumeParser):
    """
    Concrete parser implementation for DOCX resumes.
    """

    def supports(self, document: ResumeDocument) -> bool:
        """
        Return True when the document is a DOCX file.
        """

        return document.extension.lower() == "docx"

    def validate(self, document: ResumeDocument) -> None:
        """
        Validate the supplied DOCX document before parsing.
        """

        if not self.supports(document):
            raise InvalidResumeDocumentError(
                f"Unsupported file extension: {document.extension}"
            )

        if not document.file_path.exists():
            raise InvalidResumeDocumentError(
                f"File does not exist: {document.file_path}"
            )

        if not document.file_path.is_file():
            raise InvalidResumeDocumentError(
                f"Resume path is not a file: {document.file_path}"
            )

        if document.file_size <= 0:
            raise InvalidResumeDocumentError(
                "Resume file is empty."
            )

    def parse(
        self,
        document: ResumeDocument,
    ) -> ResumeParseResult:
        """
        Parse the supplied DOCX resume.
        """

        self.validate(document)

        logger.info(
            "Parsing DOCX '%s'.",
            document.file_name,
        )

        try:
            doc = Document(document.file_path)

            paragraphs: list[str] = []

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()

                if text:
                    paragraphs.append(text)

            extracted_text = "\n".join(paragraphs).strip()

            extracted_text = TextNormalizer.normalize(
                extracted_text
            )

            # Include table content because many resumes use
            # tables for skills, education, or contact information.
            table_text: list[str] = []

            for table in doc.tables:
                for row in table.rows:
                    cells: list[str] = []

                    for cell in row.cells:
                        cell_text = cell.text.strip()

                        if cell_text:
                            cells.append(cell_text)

                    if cells:
                        table_text.append(
                            " | ".join(cells)
                        )

            if table_text:
                if extracted_text:
                    extracted_text = (
                        f"{extracted_text}\n"
                        f"{chr(10).join(table_text)}"
                    )
                else:
                    extracted_text = "\n".join(
                        table_text
                    )

                extracted_text = TextNormalizer.normalize(
                    extracted_text
                )

            warnings: list[str] = []

            if not extracted_text:
                warnings.append(
                    "No text could be extracted from the DOCX document."
                )

            return ResumeParseResult(
                success=True,
                text=extracted_text,
                file_name=document.file_name,
                file_size=document.file_size,
                page_count=1,
                parser_name=self.__class__.__name__,
                warnings=warnings,
            )

        except InvalidResumeDocumentError:
            raise

        except Exception as exc:
            logger.exception(
                "Failed to parse DOCX '%s'.",
                document.file_name,
            )

            raise InvalidResumeDocumentError(
                f"Unable to parse DOCX: {document.file_name}"
            ) from exc


__all__ = ["DOCXParser"]